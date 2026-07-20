"""
事实认定书生成器

输入：结构化数据 → 生成描述 → 替换 docx 模板中的 {{truth}} → 保存

可直接作为模块被 main.py 调用，也支持独立运行（python generate_report.py）。
"""

from __future__ import annotations

import os
import sys
if getattr(sys, 'frozen', False):
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

from docx import Document
from openai import OpenAI


# =========================
# 配置区（可通过环境变量覆盖）
# =========================

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_MODEL = "deepseek-chat"
API_BASE_URL = "https://api.deepseek.com"

DEFAULT_TEMPLATE_PATH = os.path.join(BASE_DIR, "事实认定书模板.docx")


# =========================
# 1. 日期工具
# =========================

def fmt_date_cn(d: str) -> str:
    if not d:
        return ""
    parts = d.split("-")
    return f"{parts[0]}年{int(parts[1])}月{int(parts[2])}日"


def calc_days(d1: str, d2: str) -> int:
    if not d1 or not d2:
        return 0
    try:
        a = datetime.strptime(d1, "%Y-%m-%d")
        b = datetime.strptime(d2, "%Y-%m-%d")
        return (b - a).days
    except ValueError:
        return 0


# =========================
# 2. 生成事实描述
# =========================

def generate_description_template(row: dict) -> str:
    """基于模板生成事实描述"""
    r = row
    rf = fmt_date_cn(r.get("退款日期", ""))
    rt = fmt_date_cn(r.get("退回日期", ""))
    days = calc_days(r.get("退款日期", ""), r.get("退回日期", ""))

    return (
        f"未及时将国库集中支付退回资金退回国库，形成占压财政资金。"
        f"{rf}，该行预算单位零余额账户\"{r.get('账户名称', '')}\"发生1笔，"
        f"金额{r.get('金额', '')}元国库集中支付退款。"
        f"{rt}，退款经\"集中支付零余额清转待转\"账户退回国家金库咸宁市中心支库，"
        f"形成占压财政资金，占压{days}天。"
        f"检查中未发现该行从上述行为中获取违法所得。"
    )


SYSTEM_PROMPT = """你是一名专业的财政检查报告撰写助手。

## 背景
在国库集中支付制度下，零余额账户每日发生的退款应当及时清算退回国库。
当退款入账后未及时退库，即构成"占压财政资金"。

## 你的任务
根据用户提供的数据，撰写一段详细、专业的事实描述。

## 要求
- 清晰展现资金流向：何时入账→滞留→转过渡户→退回国库
- 明确指出占压天数
- 自然、专业的检查底稿语气
- 不虚构未提供的信息
- 日期用中文格式：xxxx年x月x日
- 金额用数字+元：如 18240元

直接输出描述文本，不要额外格式。"""


def generate_description_llm(row: dict, llm_config: dict = None) -> str:
    """使用大模型生成事实描述"""
    cfg = llm_config or {}
    api_base = (cfg.get("api_base") or API_BASE_URL).rstrip("/")

    # ── API Key 处理 ──
    # Ollama / 内网自部署服务不需要真实 key，传占位符即可
    # 远程 API 留空也会传占位符，调用时会返回具体认证错误，便于排查
    api_key = cfg.get("api_key") or os.environ.get("DEEPSEEK_API_KEY") or DEEPSEEK_API_KEY
    if not api_key:
        api_key = "sk-placeholder"  # 留空时传占位符，Ollama 忽略，远程 API 会报具体错误
        print("  [提示] API Key 为空，使用占位符。如连接远程服务请在配置中填写 Key")

    model = cfg.get("model") or DEEPSEEK_MODEL
    temperature = cfg.get("temperature", 0.3)
    max_tokens = cfg.get("max_tokens")
    top_p = cfg.get("top_p")

    client = OpenAI(api_key=api_key, base_url=api_base)

    user_msg = (
        f"清算日期：{row.get('清算日期', '')}\n"
        f"退款日期：{row.get('退款日期', '')}\n"
        f"退至过渡户日期：{row.get('退至过渡户日期', '（未记录）')}\n"
        f"退回日期：{row.get('退回日期', '')}\n"
        f"账户名称：{row.get('账户名称', '')}\n"
        f"金额：{row.get('金额', '')} 元\n"
        f"占压天数：{calc_days(row.get('退款日期', ''), row.get('退回日期', ''))} 天"
    )

    kwargs = dict(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_msg}
        ],
        temperature=temperature,
    )
    if max_tokens is not None:
        kwargs["max_tokens"] = int(max_tokens)
    if top_p is not None:
        kwargs["top_p"] = float(top_p)

    resp = client.chat.completions.create(**kwargs)
    return resp.choices[0].message.content


# =========================
# 3. 替换 docx 模板中的占位符
# =========================

def replace_truth_in_docx(doc: Document, truth_text: str, other_placeholders: dict = None) -> None:
    """
    遍历 docx 文档，找到所有 {{truth}} 并替换为生成的描述。
    同时替换其他 {{key}} 占位符。
    直接在原 doc 对象上修改。

    docx 模板中的占位符格式：
      {{truth}}   → 事实描述文本（由 generate_description 生成）
      {{number}}  → 文号（如 〔2026〕001号）

    替换范围覆盖 docx 的两种内容单元：
      1. 段落（Paragraphs）— 正文文本
      2. 表格（Tables）    — 表格单元格内的文本
    """
    placeholders = {"truth": truth_text}
    if other_placeholders:
        placeholders.update(other_placeholders)

    # --- 替换段落中的占位符 ---
    # 遍历每个段落 → 每个 run（同一段落中不同格式的文本片段）
    for para in doc.paragraphs:
        for run in para.runs:
            for key, value in placeholders.items():
                placeholder = "{{" + key + "}}"
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    # --- 替换表格中的占位符 ---
    # 遍历每个表格 → 每行 → 每单元格 → 每段落 → 每 run
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for key, value in placeholders.items():
                            placeholder = "{{" + key + "}}"
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)


def generate_docx(
    row: dict,
    mode: str = "template",
    output_dir: str = "output",
    output_filename: str = None,
    template_path: str = None,
    other_placeholders: dict = None,
    llm_config: dict = None,
) -> str:
    """
    从模板生成一份事实认定书 docx，返回保存路径。

    参数
    ----------
    row : dict          — 一条违规记录
    mode : str          — "template" 或 "llm"
    output_dir : str    — 输出目录
    output_filename : str — 文件名（自动生成）
    template_path : str — 模板文件路径
    other_placeholders : dict — 其他 {{key}} 占位符
    """
    if template_path is None:
        template_path = DEFAULT_TEMPLATE_PATH

    # 生成描述
    if mode == "llm":
        truth = generate_description_llm(row, llm_config=llm_config)
    else:
        truth = generate_description_template(row)

    # 加载模板
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"事实认定书模板不存在: {template_path}")
    doc = Document(template_path)

    # 替换占位符
    replace_truth_in_docx(doc, truth, other_placeholders)

    # 保存文件
    os.makedirs(output_dir, exist_ok=True)
    if not output_filename:
        year = (row.get("清算日期") or datetime.now().strftime("%Y"))[:4]
        ts = datetime.now().strftime("%H%M%S")
        output_filename = f"事实认定书_{year}{ts}.docx"

    filepath = os.path.join(output_dir, output_filename)
    doc.save(filepath)
    return os.path.abspath(filepath)


# =========================
# 4. 批量处理
# =========================

def _generate_single(row, i, total, mode, output_dir, filename, template_path, year, num, llm_config=None):
    """
    生成单份文书（供 ThreadPoolExecutor 并行调用）。
    模块级函数，线程安全。
    """
    print(f"  [{i}/{total}] 生成中...", file=sys.stderr)
    fp = generate_docx(
        row=row, mode=mode, output_dir=output_dir,
        output_filename=filename, template_path=template_path,
        other_placeholders={"number": f"〔{year}〕{num}号"},
        llm_config=llm_config,
    )
    print(f"    ✔ {fp}", file=sys.stderr)
    return fp


def batch_generate(
    rows: list[dict],
    mode: str = "template",
    output_dir: str = "output",
    template_path: str = None,
    max_workers: int = 4,
    llm_config: dict = None,
) -> list[str]:
    """
    批量生成事实认定书，返回文件路径列表。

    参数
    ----------
    rows : list[dict]      — 违规记录列表
    mode : str             — "template" 或 "llm"
    output_dir : str       — 输出目录
    template_path : str    — 模板文件路径
    max_workers : int      — 并行线程数（默认 4）。设为 1 则为顺序生成。
                             docx 生成是 I/O 密集型，多线程可显著加速。
    llm_config : dict      — LLM 配置（api_key, api_base, model, temperature, ...）
    """
    if max_workers <= 1 or len(rows) <= 1:
        # 顺序生成（原行为）
        filepaths = []
        for i, row in enumerate(rows, 1):
            print(f"  [{i}/{len(rows)}] 生成中...", file=sys.stderr)
            year = (row.get("清算日期") or datetime.now().strftime("%Y"))[:4]
            num = f"{i:03d}"
            filename = f"事实认定书_{year}{num}.docx"
            fp = generate_docx(
                row=row, mode=mode, output_dir=output_dir,
                output_filename=filename, template_path=template_path,
                other_placeholders={"number": f"〔{year}〕{num}号"},
                llm_config=llm_config,
            )
            filepaths.append(fp)
            print(f"    ✔ {fp}", file=sys.stderr)
        return filepaths

    # 并行生成（ThreadPoolExecutor — I/O 密集型任务）
    filepaths = []
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {}
        for i, row in enumerate(rows, 1):
            year = (row.get("清算日期") or datetime.now().strftime("%Y"))[:4]
            num = f"{i:03d}"
            filename = f"事实认定书_{year}{num}.docx"
            f = executor.submit(
                _generate_single,
                row=row, i=i, total=len(rows),
                mode=mode, output_dir=output_dir,
                filename=filename, template_path=template_path,
                year=year, num=num,
                llm_config=llm_config,
            )
            futures[f] = i

        for f in as_completed(futures):
            i = futures[f]
            try:
                fp = f.result()
                if fp:
                    filepaths.append((i, fp))
            except Exception as e:
                print(f"  [错误] 第 {i} 份文书生成失败: {e}", file=sys.stderr)

    # 按原始顺序排序
    filepaths.sort(key=lambda x: x[0])
    return [fp for _, fp in filepaths]


# =========================
# 5. 从结果文件生成文书
# =========================


def batch_generate_from_file(
    results_path: str,
    mode: str = "template",
    output_dir: str = "output",
    template_path: str = None,
    llm_config: dict = None,
) -> list[str]:
    """
    从已有的违规记录 Excel 文件批量生成事实认定书。

    参数
    ----------
    results_path : str       — 违规记录.xlsx 文件路径
    mode : str               — "template" 或 "llm"
    output_dir : str         — 输出目录
    template_path : str      — 模板文件路径
    llm_config : dict        — LLM 配置（api_key, api_base, model, temperature, ...）
    """
    import pandas as pd
    df = pd.read_excel(results_path)

    # 将检测结果列名映射为报告生成所需格式
    rows = []
    for _, r in df.iterrows():
        if pd.isna(r.get("退款日期")) or pd.isna(r.get("退至金库日期")):
            continue
        rows.append({
            "清算日期": r.get("清算日期", ""),
            "退款日期": r.get("退款日期", ""),
            "退至过渡户日期": r.get("退至垫款户日期", ""),
            "退回日期": r.get("退至金库日期", ""),
            "账户名称": r.get("sheet名称", ""),
            "金额": r.get("交易金额", 0),
        })

    if not rows:
        print("错误: 结果文件中没有完整的违规记录（需要同时有退款日期和退至金库日期）", file=sys.stderr)
        return []

    return batch_generate(rows, mode=mode, output_dir=output_dir, template_path=template_path, llm_config=llm_config)

# =========================
# 5. 独立运行入口（演示模式）
# =========================

def main():
    print("=" * 50)
    print("事实认定书生成器 — 独立演示模式")
    print("=" * 50)

    sample_rows = [
        {
            "清算日期": "2026-02-06",
            "退款日期": "2026-02-04",
            "退至过渡户日期": "2026-02-20",
            "退回日期": "2026-02-24",
            "账户名称": "崇阳县青少年校外活动中心",
            "金额": "18240",
        },
        {
            "清算日期": "2026-03-01",
            "退款日期": "2026-03-05",
            "退至过渡户日期": "",
            "退回日期": "2026-03-10",
            "账户名称": "通城县隽水镇财政所",
            "金额": "8500",
        },
    ]

    filepaths = batch_generate(sample_rows)
    print(f"\n共生成 {len(filepaths)} 份:")
    for fp in filepaths:
        print(f"  {fp}")


if __name__ == "__main__":
    main()